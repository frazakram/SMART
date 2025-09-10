# SMART - System for Multimodal Analysis & Response Technology
# Advanced AI Agent for Document Understanding and Multimodal Content Analysis
# Identifies content in images and allows querying them

import argparse
import base64
import glob
import io
import json
import logging
import os
import sys
import warnings
from abc import ABC, abstractmethod
from typing import Dict, List, Optional, Union

import google.generativeai as genai
import pandas as pd
import pymupdf
from langchain_text_splitters import RecursiveCharacterTextSplitter
from PIL import Image
from docx import Document
import requests


from tqdm import tqdm
import openpyxl

import csv

# Suppress warnings
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Model Abstraction Framework

class AIModel(ABC):
    """Abstract base class for AI models"""
    
    @abstractmethod
    def generate_content(self, parts: List[Dict]) -> str:
        """Generate content from the model"""
        pass
    
    @abstractmethod
    def identify_image_content(self, image: Image.Image, encoded_image: str) -> List[str]:
        """Identify content in an image"""
        pass

class GeminiModel(AIModel):
    """Implementation for Google's Gemini model"""
    
    def __init__(self, api_key: str, model_name: str = 'gemini-2.5-pro'):
        self.model_name = model_name
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel(model_name)
    
    def generate_content(self, parts: List[Dict]) -> str:
        """Generate content using Gemini"""
        try:
            response = self.model.generate_content(parts)
            return response.text
        except Exception as e:
            logger.error(f"Error calling Gemini API: {e}")
            return f"Error generating response: {str(e)}"
    
    def identify_image_content(self, encoded_image: str) -> List[str]:
        """Identify content in an image using Gemini"""
        try:
            parts = [
                {"text": """Identify what type of content this image contains. Pay special attention to graphs, charts, bar graphs, and histograms. 
                 Respond ONLY with relevant tags from this list: image, photo, text, document, handwriting, graph, chart, bar graph, histogram, diagram, 
                 table, infographic, map, drawing, screenshot.
                  Multiple tags are allowed."""},
                {"inline_data": {
                    "mime_type": "image/png",
                    "data": encoded_image
                }}
            ]
            
            response = self.model.generate_content(parts)
            
            # Extract content types
            content_types = []
            for tag in ["image", "photo", "text", "document", "handwriting", 
                        "graph", "chart", "diagram", "table", "infographic", 
                        "map", "drawing", "screenshot"]:
                if tag.lower() in response.text.lower():
                    content_types.append(tag)
            
            # Always include at least "image" as a fallback
            if not content_types:
                content_types = ["image"]
                
            return content_types
            
        except Exception as e:
            logger.error(f"Error identifying content: {e}")
            return ["image"] # Default to image

class AnthropicModel(AIModel):
    """Implementation for Anthropic's Claude model"""
    
    def __init__(self, api_key: str, model_name: str = 'claude-3-opus-20240229'):
        try:
            import anthropic
            self.client = anthropic.Anthropic(api_key=api_key)
            self.model_name = model_name
        except ImportError:
            logger.error("Anthropic SDK not installed. Run 'pip install anthropic'")
            raise
    
    def generate_content(self, parts: List[Dict]) -> str:
        """Generate content using Claude"""
        try:
            # Convert parts format from Gemini to Claude format
            system_prompt = ""
            user_message = ""
            images = []
            
            for part in parts:
                if 'text' in part:
                    if 'You are a helpful AI agent' in part['text']:  # System prompt
                        system_prompt = part['text']
                    else:
                        user_message += part['text'] + "\n\n"
                elif 'inline_data' in part and part['inline_data']['mime_type'].startswith('image/'):
                    # Add image to the message
                    image_data = part['inline_data']['data']
                    images.append({"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": image_data}})
            
            # Build the message content
            content = []
            
            # Add text
            if user_message:
                content.append({"type": "text", "text": user_message})
            
            # Add images
            for image in images:
                content.append(image)
            
            # Make API call
            response = self.client.messages.create(
                model=self.model_name,
                system=system_prompt,
                messages=[{"role": "user", "content": content}],
                max_tokens=4096
            )
            
            return response.content[0].text
        except Exception as e:
            logger.error(f"Error calling Claude API: {e}")
            return f"Error generating response: {str(e)}"
    
    def identify_image_content(self, encoded_image: str) -> List[str]:
        """Identify content in an image using Claude"""
        try:
            content = [
                {"type": "text", "text": "Identify what type of content this image contains. Respond ONLY with relevant tags from this list: image, photo, text, document, handwriting, graph, chart, diagram, table, infographic, map, drawing, screenshot. Multiple tags are allowed."},
                {"type": "image", "source": {"type": "base64", "media_type": "image/png", "data": encoded_image}}
            ]
            
            response = self.client.messages.create(
                model=self.model_name,
                messages=[{"role": "user", "content": content}],
                max_tokens=100
            )
            
            # Extract content types
            response_text = response.content[0].text
            content_types = []
            for tag in ["image", "photo", "text", "document", "handwriting", 
                        "graph", "chart", "diagram", "table", "infographic", 
                        "map", "drawing", "screenshot"]:
                if tag.lower() in response_text.lower():
                    content_types.append(tag)
            
            # Always include at least "image" as a fallback
            if not content_types:
                content_types = ["image"]
                
            return content_types
            
        except Exception as e:
            logger.error(f"Error identifying content with Claude: {e}")
            return ["image"]  # Default to image

class OpenAIModel(AIModel):
    """Implementation for OpenAI's models"""
    
    def __init__(self, api_key: str, model_name: str = 'gpt-4o'):
        try:
            import openai
            self.client = openai.OpenAI(api_key=api_key)
            self.model_name = model_name
        except ImportError:
            logger.error("OpenAI SDK not installed. Run 'pip install openai'")
            raise
    
    def generate_content(self, parts: List[Dict]) -> str:
        """Generate content using OpenAI"""
        try:
            # Convert parts format from Gemini to OpenAI format
            system_content = ""
            user_content = []
            
            for part in parts:
                if 'text' in part:
                    if 'You are a helpful AI agent' in part['text']:  # System prompt
                        system_content = part['text']
                    else:
                        user_content.append({"type": "text", "text": part['text']})
                elif 'inline_data' in part and part['inline_data']['mime_type'].startswith('image/'):
                    # Add image to the message
                    image_data = part['inline_data']['data']
                    user_content.append({
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{image_data}"}
                    })
            
            # Make API call
            messages = [
                {"role": "system", "content": system_content},
                {"role": "user", "content": user_content}
            ]
            
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=messages,
                max_tokens=4096
            )
            
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Error calling OpenAI API: {e}")
            return f"Error generating response: {str(e)}"
    
    def identify_image_content(self, image: Image.Image, encoded_image: str) -> List[str]:
        """Identify content in an image using OpenAI"""
        try:
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Identify what type of content this image contains. Respond ONLY with relevant tags from this list: image, photo, text, document, handwriting, graph, chart, diagram, table, infographic, map, drawing, screenshot. Multiple tags are allowed."},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{encoded_image}"}}
                    ]
                }
            ]
            
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=messages,
                max_tokens=100
            )
            
            # Extract content types
            response_text = response.choices[0].message.content
            content_types = []
            for tag in ["image", "photo", "text", "document", "handwriting", 
                        "graph", "chart", "diagram", "table", "infographic", 
                        "map", "drawing", "screenshot"]:
                if tag.lower() in response_text.lower():
                    content_types.append(tag)
            
            # Always include at least "image" as a fallback
            if not content_types:
                content_types = ["image"]
                
            return content_types
            
        except Exception as e:
            logger.error(f"Error identifying content with OpenAI: {e}")
            return ["image"]  # Default to image

class ModelFactory:
    """Factory class to create AI models"""
    
    @staticmethod
    def create_model(model_type: str, api_key: str, model_name: str = None) -> AIModel:
        """Create and return an AI model based on the specified type"""
        model_type = model_type.lower()
        
        if model_type == "gemini":
            default_model = "gemini-2.5-pro"
            return GeminiModel(api_key, model_name or default_model)
        elif model_type == "anthropic":
            default_model = "claude-3-opus-20240229"
            return AnthropicModel(api_key, model_name or default_model)
        elif model_type == "openai":
            default_model = "gpt-4o"
            return OpenAIModel(api_key, model_name or default_model)
        else:
            raise ValueError(f"Unsupported model type: {model_type}")

class ContentItem:
    """Class to represent an identified content item"""
    def __init__(self, 
                 content_type: str, 
                 path: str, 
                 content: Union[str, bytes], 
                 image_data: Optional[str] = None,
                 metadata: Optional[Dict] = None):
        self.type = content_type  # 'text', 'image', 'table', 'graph', 'chart'
        self.path = path
        self.content = content
        self.image_data = image_data  # Base64 encoded image
        self.metadata = metadata or {}
    
    def to_dict(self) -> Dict:
        """Convert to dictionary for API calls"""
        return {
            "type": self.type,
            "path": self.path,
            "content": self.content if isinstance(self.content, str) else None,
            "image": self.image_data,
            "metadata": self.metadata
        }

class ContentIdentifier:
    """Identifies the content of an image using an AI model."""
    def __init__(self, model):
        self.model = model

    def identify_image_content(self, image_base64: str) -> List[str]:
        """
        Identifies content types within an image using the AI model.

        Args:
            image_base64: The base64 encoded string of the image.

        Returns:
            A list of identified content tags (e.g., ['chart', 'text', 'graph']).
        """
        prompt = (
            "Analyze the following image and identify its primary content. "
            "Respond with a JSON list of keywords describing the content. "
            "For example: [\"bar chart\", \"text\", \"data table\"]. "
            "Focus on high-level categories like: 'photograph', 'document', 'screenshot', "
            "'chart', 'graph', 'diagram', 'table', 'text', 'handwriting'."
        )
        
        parts = [
            {"text": prompt},
            {"inline_data": {"mime_type": "image/png", "data": image_base64}}
        ]

        try:
            response = self.model.generate_content(parts)
            # The model might return the list in a markdown code block
            cleaned_response = response.strip().replace('`', '').replace('json', '')
            # Find the list within the response
            if "[" in cleaned_response and "]" in cleaned_response:
                list_str = cleaned_response[cleaned_response.find("["):cleaned_response.rfind("]")+1]
                identified_items = json.loads(list_str)
                if isinstance(identified_items, list):
                    return [str(item).lower() for item in identified_items]
            return []
        except (json.JSONDecodeError, Exception) as e:
            logger.error(f"Failed to identify image content: {e}")
            return []


class AIAgent:
    """SMART - System for Multimodal Analysis & Response Technology
    
    Advanced multimodal AI agent for intelligent document understanding,
    content analysis, and interactive Q&A capabilities.
    """
    
    def __init__(self, model_config: Dict = None, verification_model_config: Dict = None):
        self.items = []  # List of identified content items
        self.base_dir = "agent_data"
        self.create_directories()
        
        # Configure the primary AI model
        if model_config is None:
            # Default to Gemini if no configuration provided
            model_config = {
                "type": "gemini",
                "api_key": "GEMINI_API_KEY",  # Replace with your key
                "model_name": "gemini-2.5-pro"
            }
        
        # Create the primary AI model
        self.model = ModelFactory.create_model(
            model_type=model_config["type"],
            api_key=model_config["api_key"],
            model_name=model_config.get("model_name")
        )
        
        # Configure the verification model (if provided, otherwise use the primary model)
        if verification_model_config:
            self.verification_model = ModelFactory.create_model(
                model_type=verification_model_config["type"],
                api_key=verification_model_config["api_key"],
                model_name=verification_model_config.get("model_name")
            )
        else:
            # Default to a strong reasoning model for verification if none is provided
            self.verification_model = ModelFactory.create_model(
                model_type="gemini",
                api_key=model_config["api_key"], # Use the same API key
                model_name="gemini-1.5-flash"
            )
        
        # Content identifier uses the primary model
        self.content_identifier = ContentIdentifier(self.model)
    
    
    

    
    def create_directories(self):
        """Create necessary directories for storing processed content"""
        directories = ["images", "text", "tables", "processed", "extracted"]
        for dir_name in directories:
            os.makedirs(os.path.join(self.base_dir, dir_name), exist_ok=True)
    
    def process_input(self, input_path: str) -> List[ContentItem]:
        """Process the input (file path, URL, or directory)"""
        self.items.clear()  # Clear previous items before processing new input
        logger.info(f"Processing input: {input_path}")
        
        if input_path.startswith(('http://', 'https://')):
            return self.process_url(input_path)
        elif os.path.isdir(input_path):
            return self.process_directory(input_path)
        elif os.path.isfile(input_path):
            return self.process_file(input_path)
        else:
            logger.error(f"Invalid input path: {input_path}")
            return []
    
    def process_url(self, url: str) -> List[ContentItem]:
        """Process a URL (image or document)"""
        try:
            response = requests.get(url, stream=True)
            if response.status_code != 200:
                logger.error(f"Failed to download from URL: {url}, Status: {response.status_code}")
                return []
            
            # Determine file type from headers or URL
            content_type = response.headers.get('Content-Type', '')
            filename = url.split('/')[-1].split('?')[0]  # Extract filename from URL
            
            if not filename:
                if 'image' in content_type:
                    extension = content_type.split('/')[-1]
                    filename = f"downloaded_image.{extension}"
                elif 'pdf' in content_type:
                    filename = "downloaded_document.pdf"
                else:
                    filename = "downloaded_file"
            
            # Save the file
            filepath = os.path.join(self.base_dir, "processed", filename)
            with open(filepath, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            logger.info(f"Downloaded file to: {filepath}")
            return self.process_file(filepath)
            
        except Exception as e:
            logger.error(f"Error processing URL {url}: {e}")
            return []
    
    def process_directory(self, directory: str) -> List[ContentItem]:
        """Process all files in a directory"""
        items = []
        
        # Get all supported files
        image_extensions = ['*.jpg', '*.jpeg', '*.png', '*.bmp', '*.gif', '*.tiff']
        document_extensions = ['*.pdf', '*.docx', '*.doc', '*.txt', '*.csv', '*.xlsx', '*.xls', '*.xlsm', '*.xlsb']

        
        all_files = []
        for ext in image_extensions + document_extensions:
            all_files.extend(glob.glob(os.path.join(directory, ext)))
            all_files.extend(glob.glob(os.path.join(directory, ext.upper())))
        
        logger.info(f"Found {len(all_files)} files in directory {directory}")
        
        for file_path in tqdm(all_files, desc="Processing files"):
            items.extend(self.process_file(file_path))
        
        return items
    
    def process_file(self, file_path: str) -> List[ContentItem]:
        """Process a single file based on its type"""
        items = []
        
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Process based on file type
            if file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff']:
                items.extend(self.process_image_file(file_path))
            elif file_ext == '.pdf':
                items.extend(self.process_pdf_file(file_path))
            elif file_ext in ['.txt', '.csv']:
                items.extend(self.process_text_file(file_path))
            elif file_ext in ['.xlsx', '.xls', '.xlsm', '.xlsb']:
                items.extend(self.process_excel_file(file_path))
            elif file_ext in ['.docx', '.doc']:
                # prefer .docx; for .doc try to convert to docx first or warn
                if file_ext == '.doc':
                    logger.warning("Legacy .doc found — consider converting to .docx for better processing.")
                items.extend(self.process_docx_file(file_path))
            else:
                logger.warning(f"Unsupported file type: {file_ext}")

        
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
        
        return items
    
    def process_image_file(self, image_path: str) -> List[ContentItem]:
        """Process a single image file, creating one content item per image."""
        try:
            image_name = os.path.basename(image_path)
            with Image.open(image_path) as img:
                # Ensure image is in a consistent format (PNG)
                processed_path = os.path.join(self.base_dir, "processed", os.path.splitext(image_name)[0] + '.png')
                img.save(processed_path, 'PNG')

                # Base64 encode the image
                with open(processed_path, 'rb') as f:
                    encoded_image = base64.b64encode(f.read()).decode('utf8')

                # Identify content types in the image
                content_types = self.content_identifier.identify_image_content(encoded_image)

                # Create a single content item for the image
                item = ContentItem(
                    content_type="image",
                    path=image_path,
                    content=f"Image with identified content: {', '.join(content_types)}",
                    image_data=encoded_image,
                    metadata={"identified_content": content_types}
                )
                logger.info(f"Processed image: {image_name} - Identified: {content_types}")
                extracted_path = os.path.join(self.base_dir, "extracted", os.path.splitext(image_name)[0] + '.png')
                img.save(extracted_path, 'PNG')
                image_save_path = os.path.join(self.base_dir, "images", os.path.splitext(image_name)[0] + '.png')
                img.save(image_save_path, 'PNG')

                logger.info(f"Extracted image to: {extracted_path}")
                return [item] 
        except Exception as e:
            logger.error(f"Error processing image {image_path}: {e}")
            return []
    
    def process_pdf_file(self, pdf_path: str) -> List[ContentItem]:
        """Process a PDF file, creating one content item per page."""
        items = []
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]

        try:
            doc = pymupdf.open(pdf_path)
            logger.info(f"Processing PDF: {pdf_path} ({len(doc)} pages)")

            full_text_content = ""
            for page_num, page in enumerate(doc):
                # Save each page as a separate image
                pix = page.get_pixmap(matrix=pymupdf.Matrix(1.0, 1.0))
                image_path = os.path.join(self.base_dir, "images", f"{base_name}_page_{page_num + 1}.png")
                pix.save(image_path)

                # Encode image to base64
                with open(image_path, "rb") as img_file:
                    encoded_image = base64.b64encode(img_file.read()).decode('utf-8')

                # Create a content item for each page
                page_content = page.get_text()
                items.append(ContentItem(
                    content_type="page",
                    path=image_path,
                    content=page_content,
                    image_data=encoded_image,
                    metadata={"page_number": page_num + 1}
                ))
                full_text_content += page_content + "\n\n"

            # Save the full text content for reference, but don't create an item for it
            if full_text_content.strip():
                text_path = os.path.join(self.base_dir, "text", f"{base_name}_full.txt")
                with open(text_path, "w", encoding="utf-8") as f:
                    f.write(full_text_content)

        except Exception as e:
            logger.error(f"Error processing PDF {pdf_path}: {e}")

        return items
    
    def process_text_file(self, text_path: str) -> List[ContentItem]:
        """Process a text file"""
        items = []
        
        try:
            filename = os.path.basename(text_path)
            with open(text_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Create text splitter for chunking
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=700, 
                chunk_overlap=200, 
                length_function=len
            )
            
            # Process text chunks
            chunks = text_splitter.split_text(text)
            for i, chunk in enumerate(chunks):
                text_file_name = f"{self.base_dir}/text/{filename}_chunk_{i}.txt"
                with open(text_file_name, 'w', encoding='utf-8') as f:
                    f.write(chunk)
                
                item = ContentItem(
                    content_type="text",
                    path=text_file_name,
                    content=chunk,
                    metadata={"chunk": i}
                )
                items.append(item)
        
        except Exception as e:
            logger.error(f"Error processing text file {text_path}: {e}")
        
        return items
    def process_excel_file(self, excel_path: str) -> List[ContentItem]:
        """Process Excel files (.xlsx, .xls). Creates one ContentItem per sheet + one 'table' item per sheet."""
        items = []
        try:
            # Read all sheets as DataFrames
            sheets = pd.read_excel(excel_path, sheet_name=None, engine=None)  # pandas chooses engine
            base_name = os.path.splitext(os.path.basename(excel_path))[0]

            for sheet_name, df in sheets.items():
                # Save sheet CSV for reference
                safe_sheet_name = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in sheet_name)[:120]
                csv_path = os.path.join(self.base_dir, "tables", f"{base_name}__{safe_sheet_name}.csv")
                os.makedirs(os.path.dirname(csv_path), exist_ok=True)
                df.to_csv(csv_path, index=False, encoding='utf-8')

                # Create a table content item (include a small preview in content)
                preview = df.head(50).to_csv(index=False)
                metadata = {
                    "file": excel_path,
                    "sheet": sheet_name,
                    "rows": int(df.shape[0]),
                    "columns": int(df.shape[1]),
                    "csv_path": csv_path
                }

                item = ContentItem(
                    content_type="table",
                    path=csv_path,
                    content=preview,
                    image_data=None,
                    metadata=metadata
                )
                items.append(item)

                # Also add a text-summary item for the sheet (if needed for QA)
                text_summary = f"Sheet: {sheet_name}\nColumns: {', '.join(map(str, df.columns.tolist()))}\nRowCount: {df.shape[0]}"
                text_item_path = os.path.join(self.base_dir, "text", f"{base_name}__{safe_sheet_name}_summary.txt")
                with open(text_item_path, "w", encoding="utf-8") as f:
                    f.write(text_summary)

                items.append(ContentItem(
                    content_type="text",
                    path=text_item_path,
                    content=text_summary,
                    metadata={"sheet": sheet_name}
                ))
        except Exception as e:
            logger.error(f"Error processing excel {excel_path}: {e}")
        return items


    def process_docx_file(self, docx_path: str) -> List[ContentItem]:
        """Process Word (.docx) files. Extract paragraphs and tables as separate items."""
        items = []
        if docx is None:
            logger.error("python-docx not installed.")
            return items

        try:
            document = docx.Document(docx_path)
            base_name = os.path.splitext(os.path.basename(docx_path))[0]

            # Extract paragraphs (join into reasonable sized chunks)
            all_paras = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            joined = "\n\n".join(all_paras)
            if joined:
                # Optionally chunk using your RecursiveCharacterTextSplitter (already used for text files)
                text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=200, length_function=len)
                chunks = text_splitter.split_text(joined)
                for i, chunk in enumerate(chunks):
                    text_file_name = os.path.join(self.base_dir, "text", f"{base_name}_para_chunk_{i}.txt")
                    with open(text_file_name, "w", encoding="utf-8") as f:
                        f.write(chunk)
                    items.append(ContentItem(
                        content_type="text",
                        path=text_file_name,
                        content=chunk,
                        metadata={"chunk": i}
                    ))

            # Extract tables
            for t_index, table in enumerate(document.tables):
                rows = []
                for r in table.rows:
                    rows.append([c.text.replace("\n", " ").strip() for c in r.cells])

                # Convert to CSV string and save
                csv_path = os.path.join(self.base_dir, "tables", f"{base_name}_table_{t_index}.csv")
                os.makedirs(os.path.dirname(csv_path), exist_ok=True)
                with open(csv_path, "w", encoding="utf-8", newline='') as csvfile:
                    writer = csv.writer(csvfile)
                    for row in rows:
                        writer.writerow(row)

                # Build preview
                preview_rows = rows[:50]
                preview_csv = "\n".join([",".join([cell.replace(",", " ") for cell in r]) for r in preview_rows])
                items.append(ContentItem(
                    content_type="table",
                    path=csv_path,
                    content=preview_csv,
                    metadata={"file": docx_path, "table_index": t_index, "rows": len(rows)}
                ))

        except Exception as e:
            logger.error(f"Error processing docx {docx_path}: {e}")

        return items
    
    def query_content(self, query: str, max_items: int = 26, metrics: List[str] = None) -> str:
        """Query the identified content items with improved image context handling"""
        if not self.items:
            return "No content has been processed yet. Please provide a file, URL, or directory to analyze."
        
        # Select relevant items based on content type
        text_items = [item for item in self.items if item.type == "text"]
        table_items = [item for item in self.items if item.type == "table"]
        image_items = [item for item in self.items if item.type == "image"]
        page_items = [item for item in self.items if item.type == "page"]
        graph_items = [item for item in self.items if "graph" in item.metadata.get("identified_content", [])]
        chart_items = [item for item in self.items if "chart" in item.metadata.get("identified_content", [])]
        visualization_items = [item for item in self.items if any(vis_type in item.metadata.get("identified_content", []) 
                            for vis_type in ["graph", "chart", "bar graph", "histogram"])]

        # Select items based on the query keywords
        selected_items = []
        
        # Keywords to help with content selection
        if any(keyword in query.lower() for keyword in ["image", "picture", "photo", "show", "explain", "describe"]):
            selected_items.extend(image_items[:26])
        
        if any(keyword in query.lower() for keyword in ["text", "say", "write", "content"]):
            if text_items:
                pages_covered = set()
                for item in text_items:
                    page_num = item.metadata.get("page")
                    if page_num is not None and page_num not in pages_covered and len(pages_covered) < 5:
                        selected_items.append(item)
                        pages_covered.add(page_num)
        
        if any(keyword in query.lower() for keyword in ["table", "data", "column", "row"]):
            selected_items.extend(table_items[:26])
        
        if any(keyword in query.lower() for keyword in ["graph", "plot", "trend", "bar graph", "bar chart", "histogram","box plot","heat map","pie chart"]):
            selected_items.extend(visualization_items[:26])
        
        if any(keyword in query.lower() for keyword in ["page", "whole", "entire", "complete", "all"]):
            if page_items:
                selected_items.extend(page_items[:26])
        
        # If no specific content type was identified, use a balanced approach
        if not selected_items:
            if text_items:
                pages_covered = set()
                for item in text_items:
                    page_num = item.metadata.get("page")
                    if page_num is not None and page_num not in pages_covered and len(pages_covered) < 5:
                        selected_items.append(item)
                        pages_covered.add(page_num)
            
            if table_items:
                selected_items.extend(table_items[:26])
            
            if image_items:
                selected_items.extend(image_items[:26])
            
            if not selected_items and page_items:
                selected_items.extend(page_items[:26])
        
        # Limit to max_items
        selected_items = selected_items[:max_items]
        
        if not selected_items:
            return "I couldn't find relevant content to answer your question. Please try a different question or provide more content to analyze."
        
        # FIXED: Prepare items for the model with proper image data handling
        items_for_model = []
        for item in selected_items:
            item_dict = item.to_dict()
            # Ensure image data is properly included
            if item.image_data and not item_dict.get('image'):
                item_dict['image'] = item.image_data
            items_for_model.append(item_dict)
        
        return self.query_model(query, items_for_model, metrics)
    def verify_response(self, query: str, response: str, context_items: List[Dict], metrics: List[str]) -> Dict:
        """Verify the AI response with improved multimodal context handling"""
        results = {}
        if not metrics:
            return results

        # Build metric prompts
        metric_prompts = []
        if 'hallucination' in metrics:
            metric_prompts.append("'hallucination': {'score': <an integer from 0 to 10>, 'justification': '<a brief explanation for your score>'} (A score of 0 means a severe hallucination, while 10 means no hallucination)")
        if 'relevance' in metrics:
            metric_prompts.append("'relevance': {'score': <an integer from 0 to 10>, 'justification': '<a brief explanation for your score>'} (A score of 0 means completely irrelevant, while 10 means highly relevant)")
        if 'completeness' in metrics:
            metric_prompts.append("'completeness': {'score': <an integer from 0 to 10>, 'justification': '<a brief explanation for your score>'} (A score of 0 means very incomplete, while 10 means fully complete)")
        if 'accuracy' in metrics:
            metric_prompts.append("'accuracy': {'score': <an integer from 0 to 10>, 'justification': '<a brief explanation for your score>'} (To evaluate accuracy, examine the provided context (including images) and compare the AI's response to what you can observe. A score of 10 means the response accurately describes what is shown.)")

        system_prompt = (
            "You are an expert evaluator with access to the same multimodal context (text and images) that was used to generate the AI response. "
            "Your task is to analyze the provided evidence and return a single JSON object with your evaluation. "
            "IMPORTANT: You have access to the same images and text that the AI used to generate its response. "
            f"Evaluate for the following metrics: {', '.join(metric_prompts)}"
        )

        # FIXED: Assemble evidence with proper image handling
        evidence_parts = [
            {"text": system_prompt},
            {"text": "--- EVIDENCE BEGIN ---"},
            {"text": f"**User Query:** {query}"},
            {"text": f"**AI Response:** {response}"},
            {"text": "**Context from Documents:**"}
        ]
        
        # Add context items with explicit image handling
        has_images = False
        for idx, item in enumerate(context_items):
            if item.get('content'):
                evidence_parts.append({"text": f"Text Content {idx + 1}: {item['content']}"})
            
            # FIXED: Properly handle image data
            if item.get('image'):
                has_images = True
                evidence_parts.append({"text": f"Image {idx + 1}:"})
                evidence_parts.append({
                    "inline_data": {
                        "mime_type": "image/png",
                        "data": item['image']
                    }
                })
        
        evidence_parts.append({"text": "--- EVIDENCE END ---"})
        
        # Add context about available evidence
        if has_images:
            evidence_parts.append({"text": "Note: You have access to the same images that the AI used to generate its response. Please examine them carefully when evaluating accuracy."})

        # Get verification response
        try:
            verification_response = self.verification_model.generate_content(evidence_parts)
            logger.info(f"Verification completed with {len(context_items)} context items ({sum(1 for item in context_items if item.get('image'))} images)")
            
            # Parse verification response
            cleaned_response = verification_response.strip()
            
            # Remove markdown formatting
            if '```json' in cleaned_response:
                cleaned_response = cleaned_response.split('```json')[1].split('```')[0]
            elif '```' in cleaned_response:
                cleaned_response = cleaned_response.replace('```', '')
            
            # Extract JSON
            if '{' in cleaned_response and '}' in cleaned_response:
                json_start = cleaned_response.find('{')
                json_end = cleaned_response.rfind('}') + 1
                json_str = cleaned_response[json_start:json_end]
                verification_json = json.loads(json_str)
                
                # Process each requested metric
                for metric in metrics:
                    if metric in verification_json:
                        results[metric] = {
                            "score": verification_json[metric].get("score", 5),
                            "justification": verification_json[metric].get("justification", "No justification provided.")
                        }
                    else:
                        results[metric] = {
                            "score": 5, 
                            "justification": f"Metric '{metric}' not found in verification response."
                        }
            else:
                logger.error(f"No valid JSON found in verification response: {cleaned_response}")
                for metric in metrics:
                    results[metric] = {
                        "score": 5, 
                        "justification": "Failed to parse verification response - no valid JSON found."
                    }
                    
        except Exception as e:
            logger.error(f"Verification failed: {e}")
            for metric in metrics:
                results[metric] = {
                    "score": 5, 
                    "justification": f"Verification failed due to error: {str(e)}"
                }
        
        return results

    def query_model(self, prompt: str, content_items: List[Dict], metrics: List[str] = None) -> Dict:
        """Query the AI model with improved context handling"""
        # Prepare content parts for the model
        parts = []
        
        # System instruction
        parts.append({
            "text": "You are a helpful AI agent for content understanding and question answering. "
                "The provided text, tables, and images are relevant information retrieved to help answer the question "
                "with accuracy, completeness, clarity, and relevance. Analyze all content carefully."
        })
        
        # Add content from items with better error handling
        for idx, item in enumerate(content_items):
            try:
                if item.get('type') in ['text', 'table', 'page'] and item.get('content'):
                    parts.append({"text": f"Content {idx + 1}: {item['content']}"})
                elif item.get('image'):
                    parts.append({
                        "inline_data": {
                            "mime_type": "image/png",
                            "data": item['image']
                        }
                    })
                    logger.info(f"Added image {idx + 1} to model context")
            except Exception as e:
                logger.error(f"Error processing content item {idx}: {e}")
        
        # Add the user prompt
        parts.append({"text": prompt})
        
        # Get primary response
        response = self.model.generate_content(parts)
        
        # Verify the response if metrics are provided
        verification_results = {}
        if metrics:
            verification_results = self.verify_response(
                query=prompt,
                response=response,
                context_items=content_items,
                metrics=metrics
            )

        return {
            "response": response,
            "metrics": verification_results
        }
        # Call model API
        # return self.model.generate_content(parts)
    
    def run(self):
        """Run the agent interactively"""
        print("=" * 80)
        print("  Multimodal AI Agent for Content Understanding")
        print("  Provide a file path, URL, or directory to analyze content")
        print("=" * 80)
        
        while True:
            user_input = input("\nEnter path/URL to analyze or a question (type 'quit' to exit): ")
            
            if user_input.lower() == 'quit':
                print("Exiting the AI agent. Goodbye!")
                break
            
            # Check if this is a path/URL or a question
            if (user_input.startswith(('http://', 'https://')) or 
                os.path.exists(user_input) or
                any(ext in user_input.lower() for ext in ['.jpg', '.png', '.pdf', '/'])):
                
                # This is a path/URL to process
                print(f"Processing content from: {user_input}")
                new_items = self.process_input(user_input)
                
                if new_items:
                    self.items.extend(new_items)
                    content_types = {}
                    for item in new_items:
                        content_types[item.type] = content_types.get(item.type, 0) + 1
                    
                    print(f"\nIdentified content:")
                    for content_type, count in content_types.items():
                        print(f"- {content_type}: {count} items")
                    
                    special_content = []
                    for item in new_items:
                        if item.type == "image" and "identified_content" in item.metadata:
                            for content_type in item.metadata["identified_content"]:
                                if content_type not in ["image", "photo"]:
                                    special_content.append(content_type)
                    
                    if special_content:
                        print(f"- Special content identified: {', '.join(set(special_content))}")
                    
                    print("\nYou can now ask questions about the content.")
                else:
                    print("No content could be processed from the provided input.")
            else:
                # This is a question to answer
                if not self.items:
                    print("No content has been processed yet. Please provide a file, URL, or directory to analyze first.")
                else:
                    print("Generating response...")
                    response = self.query_content(user_input)
                    print("\nResponse:")
                    print(response)





def main():
    parser = argparse.ArgumentParser(description='Multimodal AI Agent for Content Understanding')
    parser.add_argument('--input', type=str, help='Path, URL, or directory to analyze')
    parser.add_argument('--model', type=str, default='gemini', choices=['gemini', 'anthropic', 'openai'], 
                        help='AI model to use (default: gemini)')
    parser.add_argument('--api-key', type=str, help='API key for the selected model')
    parser.add_argument('--model-name', type=str, help='Specific model name (optional)')
    parser.add_argument('--verification-model', type=str, choices=['gemini', 'anthropic', 'openai'], 
                        help='Model to use for verification (default: same as primary model)')
    parser.add_argument('--verification-api-key', type=str, help='API key for the verification model')
    parser.add_argument('--verification-model-name', type=str, help='Specific verification model name (optional)')
    
    args = parser.parse_args()
    
    # Configure the primary model
    model_config = {
        "type": args.model,
        "api_key": args.api_key or os.environ.get(f"{args.model.upper()}_API_KEY", "GEMINI_API_KEY"),
        "model_name": args.model_name
    }
    
    # Configure the verification model if specified
    verification_model_config = None
    if args.verification_model:
        verification_model_config = {
            "type": args.verification_model,
            "api_key": args.verification_api_key or os.environ.get(f"{args.verification_model.upper()}_API_KEY"),
            "model_name": args.verification_model_name
        }
    
    # Create agent with specified models
    agent = AIAgent(model_config, verification_model_config)
    
    if args.input:
        agent.process_input(args.input)
    
    agent.run()

if __name__ == "__main__":
    main()
